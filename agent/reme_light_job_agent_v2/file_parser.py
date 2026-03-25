"""
file_parser.py — 多格式文件解析工具

支持：
  - PDF (.pdf)          → 提取文字，优先 pymupdf，回退 pdfplumber
  - Word (.docx/.doc)   → 提取文字 + 表格，使用 python-docx
  - 图片 (.jpg/.png/.webp/.gif/.bmp/.tiff)
                        → base64 编码，供视觉模型分析
"""

import base64
from dataclasses import dataclass
from pathlib import Path


# ── 支持的文件类型映射 ────────────────────────────────────────────────────────
SUPPORTED_EXTENSIONS = {
    ".pdf":  "pdf",
    ".docx": "docx",
    ".doc":  "docx",
    ".jpg":  "image",
    ".jpeg": "image",
    ".png":  "image",
    ".webp": "image",
    ".gif":  "image",
    ".bmp":  "image",
    ".tiff": "image",
    ".tif":  "image",
}

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tiff", ".tif"}

MEDIA_TYPE_MAP = {
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png":  "image/png",
    ".webp": "image/webp",
    ".gif":  "image/gif",
    ".bmp":  "image/bmp",
    ".tiff": "image/tiff",
    ".tif":  "image/tiff",
}


# ── 解析结果数据类 ────────────────────────────────────────────────────────────
@dataclass
class FileParseResult:
    """统一的文件解析结果"""
    file_type: str           # "pdf" / "docx" / "image" / "unknown"
    file_name: str           # 文件名（含扩展名）
    file_path: str           # 绝对路径字符串
    text_content: str = ""   # PDF/DOCX 提取的正文文字
    image_base64: str = ""   # 图片 base64 字符串（仅图片类型）
    image_media_type: str = ""  # 图片 MIME 类型（仅图片类型）
    char_count: int = 0      # 文字字符数（图片则为 base64 长度）
    page_count: int = 0      # PDF 页数（其他类型为 0）
    error: str = ""          # 解析错误信息（成功时为空）

    @property
    def success(self) -> bool:
        """解析是否成功"""
        return not self.error

    @property
    def is_image(self) -> bool:
        return self.file_type == "image"

    @property
    def is_text_based(self) -> bool:
        return self.file_type in ("pdf", "docx")


# ── PDF 解析 ──────────────────────────────────────────────────────────────────
def parse_pdf(path: Path) -> FileParseResult:
    """
    提取 PDF 文本内容。
    优先使用 pymupdf（中文/排版支持更好），回退到 pdfplumber。
    """
    result = FileParseResult(
        file_type="pdf",
        file_name=path.name,
        file_path=str(path.absolute()),
    )

    # ── 方案 1：pymupdf (fitz) ──────────────────────────────────────────────
    try:
        import fitz  # pip install pymupdf
        doc = fitz.open(str(path))
        result.page_count = len(doc)
        texts = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")  # 纯文字模式，保留换行
            if text.strip():
                texts.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
        doc.close()
        result.text_content = "\n\n".join(texts).strip()
        result.char_count = len(result.text_content)
        return result
    except ImportError:
        pass  # 未安装，尝试下一个方案
    except Exception as e:
        result.error = f"pymupdf 解析失败: {e}"
        # 不立即返回，尝试 pdfplumber

    # ── 方案 2：pdfplumber ──────────────────────────────────────────────────
    try:
        import pdfplumber  # pip install pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            result.page_count = len(pdf.pages)
            texts = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    texts.append(f"--- 第 {i + 1} 页 ---\n{text}")

                # 提取表格（简历常见）
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = " | ".join(
                            (cell or "").strip() for cell in row if cell
                        )
                        if row_text:
                            texts.append(row_text)

        result.text_content = "\n\n".join(texts).strip()
        result.char_count = len(result.text_content)
        result.error = ""  # 清除之前的错误
        return result
    except ImportError:
        pass
    except Exception as e:
        result.error = f"pdfplumber 解析失败: {e}"

    # 两个方案都失败
    if not result.error:
        result.error = (
            "未找到 PDF 解析库，请安装：\n"
            "  pip install pymupdf\n"
            "或者\n"
            "  pip install pdfplumber"
        )
    return result


# ── DOCX 解析 ─────────────────────────────────────────────────────────────────
def parse_docx(path: Path) -> FileParseResult:
    """
    提取 Word (.docx) 文本内容，包含正文段落和表格。
    使用 python-docx。
    """
    result = FileParseResult(
        file_type="docx",
        file_name=path.name,
        file_path=str(path.absolute()),
    )

    try:
        from docx import Document  # pip install python-docx
        doc = Document(str(path))
        parts = []

        # 提取正文段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 保留标题样式标记，便于 LLM 识别结构
                if para.style.name.startswith("Heading"):
                    parts.append(f"\n### {text}")
                else:
                    parts.append(text)

        # 提取表格（简历、对比表格常见）
        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells]
                # 去重（相邻合并单元格会重复）
                seen = []
                for t in row_texts:
                    if t and (not seen or seen[-1] != t):
                        seen.append(t)
                if seen:
                    table_lines.append(" | ".join(seen))
            if table_lines:
                parts.append("\n[表格]\n" + "\n".join(table_lines))

        result.text_content = "\n".join(parts).strip()
        result.char_count = len(result.text_content)
        return result

    except ImportError:
        result.error = "未找到 python-docx，请安装：pip install python-docx"
    except Exception as e:
        result.error = f"DOCX 解析失败: {e}"

    return result


# ── 图片编码 ──────────────────────────────────────────────────────────────────
def encode_image(path: Path) -> FileParseResult:
    """
    将图片文件编码为 base64 字符串，供视觉模型 API 使用。
    可选：用 Pillow 做基础校验并读取尺寸信息。
    """
    ext = path.suffix.lower()
    media_type = MEDIA_TYPE_MAP.get(ext, "image/jpeg")

    result = FileParseResult(
        file_type="image",
        file_name=path.name,
        file_path=str(path.absolute()),
        image_media_type=media_type,
    )

    try:
        # 用 Pillow 读取基础信息（可选，不影响主流程）
        try:
            from PIL import Image  # pip install Pillow
            with Image.open(str(path)) as img:
                w, h = img.size
                fmt = img.format or ext.lstrip(".")
                result.text_content = (
                    f"[图片元信息] 文件: {path.name} | "
                    f"尺寸: {w}x{h}px | 格式: {fmt.upper()}"
                )
        except ImportError:
            result.text_content = f"[图片] {path.name}"
        except Exception:
            result.text_content = f"[图片] {path.name}"

        # base64 编码（核心步骤）
        with open(str(path), "rb") as f:
            raw = f.read()
        result.image_base64 = base64.b64encode(raw).decode("utf-8")
        result.char_count = len(raw)  # 原始字节数（更直观）
        return result

    except Exception as e:
        result.error = f"图片处理失败: {e}"
        return result


# ── 统一入口 ──────────────────────────────────────────────────────────────────
def parse_file(path: Path) -> FileParseResult:
    """
    统一文件解析入口。根据扩展名自动路由到对应解析器。

    Args:
        path: 文件路径（Path 对象）

    Returns:
        FileParseResult，通过 .success 判断是否解析成功
    """
    if not path.exists():
        return FileParseResult(
            file_type="unknown",
            file_name=path.name,
            file_path=str(path.absolute()),
            error=f"文件不存在: {path}",
        )

    if not path.is_file():
        return FileParseResult(
            file_type="unknown",
            file_name=path.name,
            file_path=str(path.absolute()),
            error=f"路径不是文件: {path}",
        )

    ext = path.suffix.lower()
    file_type = SUPPORTED_EXTENSIONS.get(ext, "unknown")

    if file_type == "pdf":
        return parse_pdf(path)
    elif file_type == "docx":
        return parse_docx(path)
    elif file_type == "image":
        return encode_image(path)
    else:
        return FileParseResult(
            file_type="unknown",
            file_name=path.name,
            file_path=str(path.absolute()),
            error=(
                f"不支持的文件类型: '{ext}'\n"
                f"支持的格式: PDF (.pdf) | Word (.docx/.doc) | "
                f"图片 (.jpg/.png/.webp/.gif/.bmp/.tiff)"
            ),
        )


# ── 工具函数 ──────────────────────────────────────────────────────────────────
def guess_content_role(text: str) -> str:
    """
    粗略判断文本是简历还是 JD，辅助自动提示。
    返回: "resume" / "jd" / "unknown"
    """
    text_lower = text.lower()
    jd_keywords   = ["岗位职责", "任职要求", "job description", "responsibilities",
                     "requirements", "我们提供", "薪资", "招聘", "职位"]
    res_keywords   = ["工作经历", "教育背景", "个人简介", "技能", "项目经验",
                     "自我评价", "联系方式", "求职意向", "毕业院校"]

    jd_score  = sum(1 for kw in jd_keywords  if kw in text_lower)
    res_score = sum(1 for kw in res_keywords if kw in text_lower)

    if jd_score > res_score:
        return "jd"
    elif res_score > jd_score:
        return "resume"
    return "unknown"


def get_supported_formats() -> str:
    """返回支持格式的可读描述"""
    return "PDF (.pdf) | Word (.docx/.doc) | 图片 (.jpg / .png / .webp / .gif / .bmp)"